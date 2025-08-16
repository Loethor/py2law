from app.services.doc_generator import create_docx, populate_template
from docx import Document


def test_create_docx(tmp_path):
    """
    Test the create_docx function to ensure it creates a valid .docx file.
    """
    # Define the path for the test .docx file
    test_file_path = tmp_path / "test_document.docx"

    # Call the create_docx function
    success = create_docx(str(test_file_path))

    # Assert that the function returns True
    assert success, "create_docx should return True on success"

    # Assert that the file was created
    assert test_file_path.exists(), "The .docx file should be created"

    # Assert that the file is not empty
    assert test_file_path.stat().st_size > 0, "The .docx file should not be empty"

    # Verify the content of the .docx file
    document = Document(test_file_path)
    paragraphs = [p.text for p in document.paragraphs]

    # Check for the title
    assert (
        "Documento de Ejemplo" in paragraphs
    ), "The document should contain the title 'Documento de Ejemplo'"

    # Check for sample content
    assert (
        "Este es un documento de ejemplo generado automáticamente." in paragraphs
    ), "The document should contain the sample content"
    assert (
        "Puedes personalizar este contenido según tus necesidades." in paragraphs
    ), "The document should contain the customizable content"

    # Check for styled text
    assert any(
        "Este texto tiene un tamaño de fuente personalizado." in p for p in paragraphs
    ), "The document should contain styled text"


def test_populate_template(tmp_path):
    """
    Test the populate_template function to ensure it correctly replaces placeholders in the template.
    """
    # Create a sample template file
    template_path = tmp_path / "template.docx"
    document = Document()
    document.add_paragraph("Nombre: {{name}}")
    document.add_paragraph("Fecha: {{date}}")
    document.add_paragraph("Contenido: {{content}}")
    document.save(template_path)

    # Define the path for the populated .docx file
    output_path = tmp_path / "populated_document.docx"

    # Define the context for placeholder replacement
    context = {
        "name": "Juan Pérez",
        "date": "01/01/2023",
        "content": "Este es el contenido de la reclamación previa.",
    }

    # Call the populate_template function
    success = populate_template(str(template_path), str(output_path), context)

    # Assert that the function returns True
    assert success, "populate_template should return True on success"

    # Assert that the populated file was created
    assert output_path.exists(), "The populated .docx file should be created"

    # Assert that the file is not empty
    assert (
        output_path.stat().st_size > 0
    ), "The populated .docx file should not be empty"

    # Verify the content of the populated .docx file
    populated_document = Document(output_path)
    paragraphs = [p.text for p in populated_document.paragraphs]

    # Check that placeholders were replaced with the correct values
    assert (
        "Nombre: Juan Pérez" in paragraphs
    ), "The placeholder {{name}} should be replaced with 'Juan Pérez'"
    assert (
        "Fecha: 01/01/2023" in paragraphs
    ), "The placeholder {{date}} should be replaced with '01/01/2023'"
    assert (
        "Contenido: Este es el contenido de la reclamación previa." in paragraphs
    ), "The placeholder {{content}} should be replaced with the provided content"
