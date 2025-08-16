from docx import Document
from docx.shared import Pt
import os


def create_docx(file_path):
    """
    Creates a .docx file with sample content.

    Args:
        file_path (str): The full path where the .docx file will be saved.
    """
    try:
        # Create a new Word document
        document = Document()

        # Add a title
        document.add_heading("Documento de Ejemplo", level=1)

        # Add some sample content
        document.add_paragraph(
            "Este es un documento de ejemplo generado automáticamente."
        )
        document.add_paragraph(
            "Puedes personalizar este contenido según tus necesidades."
        )

        # Add a styled paragraph
        paragraph = document.add_paragraph("Texto con estilo: ")
        run = paragraph.add_run("Este texto tiene un tamaño de fuente personalizado.")
        run.font.size = Pt(12)

        # Save the document to the specified path
        document.save(file_path)
        return True
    except Exception as e:
        print(f"Error saving document: {e}")
        return False


def populate_template(template_path, output_path, context):
    """
    Populates a .docx template with the given context.

    Args:
        template_path (str): Path to the .docx template.
        output_path (str): Path to save the populated .docx file.
        context (dict): Dictionary with placeholder keys and their replacement values.

    Returns:
        bool: True if the document is successfully populated and saved, False otherwise.
    """
    try:
        # Load the template
        document = Document(template_path)

        # Replace placeholders with actual values
        for paragraph in document.paragraphs:
            for key, value in context.items():
                if f"{{{{{key}}}}}" in paragraph.text:  # Check for placeholder
                    paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", value)

        # Save the populated document
        document.save(output_path)
        return True
    except Exception as e:
        print(f"Error populating template: {e}")
        return False


def generate_jubilaciones_document(data, output_path):
    """
    Generates a 'Reclamación Previa a la Jubilación' document by populating the template.

    Args:
        data (dict): Dictionary containing the values for the placeholders in the template.
        output_path (str): Path to save the generated document.

    Returns:
        bool: True if the document is successfully generated, False otherwise.
    """
    try:
        # Path to the template
        template_path = os.path.join(
            "app", "templates", "Reclamación Previa a la Jubilación.docx"
        )

        # Load the template
        document = Document(template_path)

        # Replace placeholders in the document
        for paragraph in document.paragraphs:
            for key, value in data.items():
                placeholder = f"{{{{{key}}}}}"  # Format the placeholder as {{key}}
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)

        # Save the populated document
        document.save(output_path)
        return True
    except Exception as e:
        print(f"Error generating document: {e}")
        return False
